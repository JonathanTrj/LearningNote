#coding:utf-8

from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

table = doc.add_table(3,2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i in range(0, 3):
	for j in range(0,2):
		cell = table.cell(i,j)
		ph = cell.add_paragraph()
		run = ph.add_run()
		run.add_picture("wooo.jpg", width=Cm(7.5))

doc.save("b.docx")
